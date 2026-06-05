"""
Генерация Word-документа БРС из ответа Яндекс Формы.

Берёт «сырой» JSON, который присылает интеграция Яндекс Форм, и заполняет
по нему шаблон Таблицы 1.1 (data/*.docx), сохраняя оформление 1-в-1.

Маппинг полей зафиксирован по стабильному question.id (см. captures/ и
контрольный тест от 2026-06-05). Пустые ответы в JSON Яндекса ОТСУТСТВУЮТ —
их трактуем как пустые значения / 0.
"""

from __future__ import annotations

import re
from copy import deepcopy
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

TABLE_FONT_NAME = "Times New Roman"
TABLE_FONT_SIZE = 10

# ---- Идентификаторы вопросов формы (question.id) ----
ID_FAMILIYA = 125700719
ID_IMYA = 125700737
ID_OTCHESTVO = 125700764
ID_DISCIPLINA = 125700776
ID_GROUP = 125701457        # answer_group: массив строк таблицы
ID_EKZAMEN = 125704063      # промежуточная аттестация (экзамен/зачёт), одно число
ID_TEMA = 125701501
ID_RAZDEL = 125701529
ID_INDIKATORY = 125701531

# Числовые колонки баллов внутри строки группы — в порядке колонок 3..10 таблицы.
SCORE_IDS = [
    125702036,  # Устный/письменный ответ
    125702037,  # Тест
    125702039,  # Коллоквиум
    125702043,  # Контрольная работа
    125702050,  # Отчёт по лабораторной работе
    125702051,  # Доклад
    125702052,  # Реферат/эссе
    125702053,  # Защита курсового проекта/работы
]

# Структура шаблона data/*.docx (таблица 0):
TABLE_DATA_FROM = 3   # первая строка-пример (тема)
TABLE_DATA_TO = 8     # последняя строка-пример (включительно)
INTRO_PARA_INDEX = 2  # абзац с названием дисциплины


@dataclass
class Row:
    tema: str = ""
    razdel: str = ""
    indikatory: str = ""
    scores: list[int] = field(default_factory=lambda: [0] * len(SCORE_IDS))

    @property
    def bally_temy(self) -> int:
        return sum(self.scores)


@dataclass
class Submission:
    familiya: str = ""
    imya: str = ""
    otchestvo: str = ""
    disciplina: str = ""
    ekzamen: int = 0
    rows: list[Row] = field(default_factory=list)

    @property
    def fio_parts(self) -> list[str]:
        return [p for p in (self.familiya, self.imya, self.otchestvo) if p]


# --------------------------------------------------------------------------- #
# Парсинг ответа формы
# --------------------------------------------------------------------------- #

def _index_by_id(data: dict) -> dict:
    """Преобразует {slug: {value, question:{id}}} -> {question_id: value}."""
    out: dict = {}
    for entry in data.values():
        if isinstance(entry, dict):
            qid = entry.get("question", {}).get("id")
            if qid is not None:
                out[qid] = entry.get("value")
    return out


def _as_int(value) -> int:
    try:
        return int(value)
    except (TypeError, ValueError):
        return 0


def _as_text(value) -> str:
    return "" if value is None else str(value)


def _cap_first(text: str) -> str:
    """Делает первую букву заглавной, остальной текст не трогает."""
    text = text.strip()
    return text[:1].upper() + text[1:] if text else text


def _normalize_indicators(text: str) -> str:
    """
    Приводит индикаторы к виду «ПК-2 ОПК-6.1»: верхний регистр,
    разделение пробелом (в форме могут быть через запятую или пробел).
    """
    tokens = [t for t in re.split(r"[,\s]+", text.strip()) if t]
    return " ".join(t.upper() for t in tokens)


def parse_submission(body: dict) -> Submission:
    """Принимает распарсенное тело запроса Яндекс Формы и возвращает Submission."""
    top = _index_by_id(body.get("answer", {}).get("data", {}))

    sub = Submission(
        familiya=_as_text(top.get(ID_FAMILIYA)),
        imya=_as_text(top.get(ID_IMYA)),
        otchestvo=_as_text(top.get(ID_OTCHESTVO)),
        disciplina=_as_text(top.get(ID_DISCIPLINA)),
        ekzamen=_as_int(top.get(ID_EKZAMEN)),
    )

    group = top.get(ID_GROUP) or []
    if isinstance(group, list):
        for raw in group:
            if not isinstance(raw, dict):
                continue
            ri = _index_by_id(raw)
            sub.rows.append(
                Row(
                    tema=_as_text(ri.get(ID_TEMA)),
                    razdel=_as_text(ri.get(ID_RAZDEL)),
                    indikatory=_as_text(ri.get(ID_INDIKATORY)),
                    scores=[_as_int(ri.get(cid)) for cid in SCORE_IDS],
                )
            )
    return sub


# --------------------------------------------------------------------------- #
# Заполнение шаблона Word
# --------------------------------------------------------------------------- #

def _set_cell(cell, value) -> None:
    """Записывает текст в ячейку, сохраняя форматирование первого run'а."""
    text = "" if value in (None, "", 0) else str(value)
    p = cell.paragraphs[0]
    if p.runs:
        p.runs[0].text = text
        for r in list(p.runs[1:]):
            r._r.getparent().remove(r._r)
    else:
        p.add_run(text)
    for extra in list(cell.paragraphs[1:]):
        extra._p.getparent().remove(extra._p)


def _apply_table_font(table, name: str = TABLE_FONT_NAME, size: int = TABLE_FONT_SIZE) -> None:
    """
    Проставляет шрифт и размер ВСЕМУ тексту внутри таблицы.

    Идём по XML на уровне каждого run'а (w:r), чтобы не пропустить ни одного
    (обход через para.runs пропускал часть). size — в пунктах.
    """
    half_points = str(int(size * 2))  # w:sz задаётся в полупунктах
    for r in table._tbl.iter(qn("w:r")):
        rpr = r.find(qn("w:rPr"))
        if rpr is None:
            rpr = r.makeelement(qn("w:rPr"), {})
            r.insert(0, rpr)

        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = rpr.makeelement(qn("w:rFonts"), {})
            rpr.insert(0, rfonts)  # rFonts должен идти первым в rPr
        # ascii/hAnsi — латиница, cs/eastAsia — кириллица и пр.
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rfonts.set(qn(attr), name)

        for tag in ("w:sz", "w:szCs"):
            el = rpr.find(qn(tag))
            if el is None:
                el = rpr.makeelement(qn(tag), {})
                rpr.append(el)
            el.set(qn("w:val"), half_points)


def _replace_discipline(doc, disciplina: str) -> None:
    """Подставляет название дисциплины в подчёркнутый run вводного абзаца."""
    if not disciplina:
        return
    try:
        para = doc.paragraphs[INTRO_PARA_INDEX]
    except IndexError:
        return
    value = _cap_first(disciplina)
    for run in para.runs:
        if run.underline:
            run.text = value
            return
    # запасной вариант — известная позиция run'а
    if len(para.runs) > 2:
        para.runs[2].text = value


def _sanitize_filename(name: str) -> str:
    cleaned = re.sub(r'[\\/:*?"<>|]+', "", name).strip().strip("_").rstrip(" .")
    return cleaned or "no_name"


def _initials(imya: str, otchestvo: str) -> str:
    """Имя и отчество -> инициалы вида «И.О.»."""
    return "".join(f"{p[0].upper()}." for p in (imya, otchestvo) if p)


def find_template(base_dir: Path) -> Path:
    """Находит .docx-шаблон в папке data/."""
    candidates = sorted((base_dir / "data").glob("*.docx"))
    if not candidates:
        raise FileNotFoundError("Шаблон .docx не найден в папке data/")
    return candidates[0]


def generate_docx(body: dict, base_dir: Path, output_dir: Path) -> Path:
    """
    Генерирует Word-документ БРС из тела запроса Яндекс Формы.
    Возвращает путь к сохранённому файлу.
    """
    sub = parse_submission(body)
    template = find_template(base_dir)
    doc = Document(str(template))

    _replace_discipline(doc, sub.disciplina)

    table = doc.tables[0]
    template_tr = deepcopy(table.rows[TABLE_DATA_FROM]._tr)
    r9_tr = table.rows[9]._tr  # строка «Промежуточная аттестация»

    # Удаляем строки-примеры R3..R8.
    for i in range(TABLE_DATA_TO, TABLE_DATA_FROM - 1, -1):
        tr = table.rows[i]._tr
        tr.getparent().remove(tr)

    # Вставляем по одной строке на каждую тему из формы.
    for _ in sub.rows:
        r9_tr.addprevious(deepcopy(template_tr))

    # Баллы раздела = сумма баллов тем с одинаковым названием раздела.
    razdel_sums: dict[str, int] = {}
    for r in sub.rows:
        razdel_sums[r.razdel] = razdel_sums.get(r.razdel, 0) + r.bally_temy

    n = len(sub.rows)
    for k, row in enumerate(sub.rows):
        cells = table.rows[TABLE_DATA_FROM + k].cells
        _set_cell(cells[0], _normalize_indicators(row.indikatory))
        _set_cell(cells[1], _cap_first(row.razdel))
        _set_cell(cells[2], _cap_first(row.tema))
        for j, val in enumerate(row.scores):
            _set_cell(cells[3 + j], val)
        _set_cell(cells[11], "")            # экзамен в строках темы — пусто
        _set_cell(cells[12], row.bally_temy)
        _set_cell(cells[13], razdel_sums.get(row.razdel, 0))

    # Строки «Промежуточная аттестация» (R9) и «ИТОГО» (R10) — индексы сдвинулись.
    r9 = table.rows[TABLE_DATA_FROM + n]
    r10 = table.rows[TABLE_DATA_FROM + n + 1]

    _set_cell(r9.cells[11], sub.ekzamen)

    col_sums = [sum(r.scores[j] for r in sub.rows) for j in range(len(SCORE_IDS))]
    total_current = sum(r.bally_temy for r in sub.rows)
    for j, cs in enumerate(col_sums):
        _set_cell(r10.cells[3 + j], cs)
    _set_cell(r10.cells[11], sub.ekzamen)
    _set_cell(r10.cells[12], total_current)
    _set_cell(r10.cells[13], total_current + sub.ekzamen)

    # Единый шрифт для всей таблицы.
    _apply_table_font(table)

    # Имя файла: БРС_Дисциплина_Фамилия_И.О..docx
    segments = ["БРС"]
    if sub.disciplina:
        segments.append(_cap_first(sub.disciplina))
    if sub.familiya:
        segments.append(_cap_first(sub.familiya))
    initials = _initials(sub.imya, sub.otchestvo)
    if initials:
        segments.append(initials)
    fname = "_".join(segments)
    output_dir.mkdir(exist_ok=True)
    out_path = output_dir / (_sanitize_filename(fname) + ".docx")
    doc.save(str(out_path))
    return out_path


if __name__ == "__main__":
    import json
    import sys

    base = Path(__file__).parent
    arg = Path(sys.argv[1]) if len(sys.argv) > 1 else None
    if arg is None:
        caps = sorted((base / "captures").glob("request_*.json"))
        arg = caps[-1] if caps else None
    if arg is None:
        print("Нет файлов в captures/")
        raise SystemExit(1)

    cap = json.loads(arg.read_text(encoding="utf-8"))
    out = generate_docx(cap["body_parsed"], base, base / "output")
    print("Сохранён:", out)
